#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import json
import os
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from PIL import Image, ImageChops
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from design_document_diagrams import (
    build_responsibility_diagram,
    build_status_diagram,
    build_workflow_diagram,
)

ACCENT = "315EFB"
ACCENT_LIGHT = "E8EEFF"
WHITE = "FFFFFF"
LIGHT_ROW = "F4F6FA"
ALLOWED_VISUAL_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp"}
ACTIVE_STATUSES = {"ACTIVE", "CURRENT"}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build human-readable DOCX/PDF game design documents from structured JSON sources."
    )
    parser.add_argument("--registry", required=True, help="DESIGN_DOCUMENT_REGISTRY.json")
    parser.add_argument("--only", action="append", default=[], help="Build only the given document_id; repeatable")
    parser.add_argument("--source-commit", default=os.environ.get("GITHUB_SHA", ""))
    parser.add_argument(
        "--human-visual-review",
        default="NOT_RUN",
        choices=["NOT_RUN", "PASSED", "FAILED"],
    )
    return parser.parse_args()


def load_json(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError(f"JSON root must be an object: {path}")
    return data


def digest(path: Path) -> str:
    hasher = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def resolve(base: Path, value: str) -> Path:
    path = Path(value)
    return path if path.is_absolute() else (base / path).resolve()


def set_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd")) or OxmlElement("w:shd")
    if node.getparent() is None:
        props.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if margins.getparent() is None:
        props.append(margins)
    for name, value in (("top", 80), ("start", 95), ("bottom", 80), ("end", 95)):
        node = margins.find(qn(f"w:{name}")) or OxmlElement(f"w:{name}")
        if node.getparent() is None:
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_run(run, size: float = 9, bold: bool = False, color: tuple[int, int, int] | None = None) -> None:
    run.font.name = "NanumGothic"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "NanumGothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def configure_styles(doc: Document) -> None:
    for style_name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "NanumGothic"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "NanumGothic")
    doc.styles["Normal"].font.size = Pt(9)
    doc.styles["Title"].font.size = Pt(30)
    doc.styles["Title"].font.bold = True
    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 1"].font.color.rgb = RGBColor(49, 94, 251)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 2"].font.bold = True


def _value(value: Any) -> str:
    if isinstance(value, list):
        return "\n".join(str(item) for item in value) if value else "-"
    if isinstance(value, dict):
        return " / ".join(f"{key}: {_value(item)}" for key, item in value.items()) if value else "-"
    if value in (None, ""):
        return "-"
    if isinstance(value, bool):
        return "예" if value else "아니오"
    return str(value)


def add_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, ACCENT)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(cell.paragraphs[0].add_run(str(header)), 7.6, True, (255, 255, 255))
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            set_cell_shading(cell, WHITE if row_index % 2 == 0 else LIGHT_ROW)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            style_run(cell.paragraphs[0].add_run(_value(value)), 7.1)
    doc.add_paragraph()


def add_bullets(doc: Document, values: list[Any]) -> None:
    for value in values or ["[작성 필요]"]:
        paragraph = doc.add_paragraph(style="List Bullet")
        style_run(paragraph.add_run(_value(value)), 9)


def add_metadata_table(doc: Document, document: dict[str, Any], source_hash: str, source_commit: str) -> None:
    metadata = document.get("metadata", {})
    pairs = [
        ("프로젝트", document.get("project", "[프로젝트명]")),
        ("분야", document.get("discipline", "[분야]")),
        ("문서 종류", document.get("document_kind", "discipline-bible")),
        ("책임", document.get("owner", "[책임자]")),
        ("상태", document.get("status", "DRAFT")),
        ("마지막 검토", metadata.get("last_reviewed", "[작성 필요]")),
        ("기준 커밋", source_commit or metadata.get("source_commit", "[작성 필요]")),
        ("JSON SHA-256", source_hash[:20] + "…"),
        ("현재 제품 게이트", metadata.get("current_product_gate", "[작성 필요]")),
        ("현재 작업 게이트", metadata.get("current_work_gate", "[작성 필요]")),
        ("다음 검토 트리거", metadata.get("next_review_trigger", "[작성 필요]")),
        ("사람용 형식", "DOCX + PDF + 다이어그램·승인 이미지"),
    ]
    table = doc.add_table(rows=(len(pairs) + 1) // 2, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, (label, value) in enumerate(pairs):
        row, pair = divmod(index, 2)
        label_cell = table.rows[row].cells[pair * 2]
        value_cell = table.rows[row].cells[pair * 2 + 1]
        set_cell_shading(label_cell, ACCENT_LIGHT)
        set_cell_margins(label_cell)
        set_cell_margins(value_cell)
        style_run(label_cell.paragraphs[0].add_run(label), 8.5, True)
        style_run(value_cell.paragraphs[0].add_run(_value(value)), 8.5)
    doc.add_paragraph()


def add_header_footer(doc: Document, title: str, source_hash: str) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = title
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            style_run(run, 7.5, color=(102, 112, 133))
        footer = section.footer.paragraphs[0]
        footer.text = f"자동 생성 사람용 파생본 | JSON SHA-256 {source_hash[:16]}…"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            style_run(run, 7, color=(102, 112, 133))


def add_image(doc: Document, path: Path, caption: str, width: float = 6.9) -> None:
    doc.add_picture(str(path), width=Inches(width))
    paragraph = doc.add_paragraph(caption)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        style_run(run, 8, color=(102, 112, 133))


def normalize_table_rows(items: list[dict[str, Any]], fields: list[str]) -> list[list[str]]:
    return [[_value(item.get(field)) for field in fields] for item in items]


def build_docx(
    document: dict[str, Any],
    source_hash: str,
    source_commit: str,
    generated_diagrams: dict[str, Path],
    approved_visuals: list[tuple[dict[str, Any], Path]],
    output: Path,
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(1.65)
    configure_styles(doc)
    title_text = str(document.get("title", "[기획서 제목]"))
    add_header_footer(doc, title_text, source_hash)

    title = doc.add_paragraph(style="Title")
    title.add_run(title_text)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("AI용 구조화 JSON에서 자동 생성한 사람용 최신 통합본")
    add_metadata_table(doc, document, source_hash, source_commit)

    overview = document.get("overview", {})
    doc.add_heading("1. 한눈에 보기", level=1)
    add_table(doc, ["항목", "현재 기준"], [
        ["분야 목적", overview.get("purpose")],
        ["플레이어·사용자 가치", overview.get("player_value")],
        ["현재 목표", overview.get("current_goal")],
        ["현재 요약", overview.get("summary")],
        ["가장 큰 위험", overview.get("largest_risk")],
        ["다음 작업", overview.get("next_action")],
    ])
    add_image(doc, generated_diagrams["status"], "그림 1. 확정·구현·검증·확인 필요·보류 상태 분리")

    doc.add_heading("2. 목표와 Quality Bar", level=1)
    add_bullets(doc, document.get("goals", []))
    quality = document.get("quality_bar", [])
    add_table(doc, ["기준", "합격 조건", "증거", "상태"], normalize_table_rows(quality, ["criterion", "pass_condition", "evidence", "status"]) or [["[작성 필요]", "-", "-", "미확정"]])
    doc.add_heading("금지 방향", level=2)
    add_bullets(doc, document.get("prohibited_directions", []))

    doc.add_heading("3. 책임과 협업 경계", level=1)
    add_image(doc, generated_diagrams["responsibility"], "그림 2. 분야 책임·협업 경계")
    responsibilities = document.get("responsibilities", {})
    doc.add_heading("이 분야가 소유", level=2)
    add_bullets(doc, responsibilities.get("owns", []))
    doc.add_heading("이 분야가 소유하지 않음", level=2)
    add_bullets(doc, responsibilities.get("does_not_own", []))
    interfaces = responsibilities.get("interfaces", [])
    add_table(doc, ["상대 분야", "받는 입력", "제공하는 출력", "함께 갱신할 원본"], normalize_table_rows(interfaces, ["discipline", "receives", "provides", "update_together"]) or [["[작성 필요]", "-", "-", "-"]])

    doc.add_heading("4. 분야 전체 작업 과정", level=1)
    add_image(doc, generated_diagrams["workflow"], "그림 3. 입력부터 다음 게이트까지의 전체 작업 흐름")
    workflow = document.get("workflow", [])
    add_table(doc, ["단계", "입력", "판단·작업", "산출물", "실제 경로", "실패·폴백", "검증", "다음 게이트"], normalize_table_rows(workflow, ["step", "inputs", "decision_or_work", "outputs", "paths", "fallback", "validation", "next_gate"]) or [["[작성 필요]", "-", "-", "-", "-", "-", "-", "-"]])

    doc.add_heading("5. 개발 게이트", level=1)
    gates = document.get("development_gates", {})
    add_table(doc, ["작업 게이트", "진입 조건", "종료 기준", "증거", "상태"], normalize_table_rows(gates.get("work", []), ["gate", "entry", "exit", "evidence", "status"]) or [["[작성 필요]", "-", "-", "-", "-"]])
    add_table(doc, ["제품 단계", "이 분야가 증명할 것", "Quality Bar", "증거", "상태"], normalize_table_rows(gates.get("product", []), ["gate", "proof", "quality_bar", "evidence", "status"]) or [["[작성 필요]", "-", "-", "-", "-"]])

    doc.add_heading("6. 프로젝트 스킬", level=1)
    skills = document.get("skills", {})
    add_table(doc, ["구분", "스킬·경로"], [
        ["Foundation", skills.get("foundation")],
        ["분야 진입 스킬", skills.get("discipline_entry")],
        ["하위 스킬", skills.get("subskills")],
        ["Learning Log", skills.get("learning_log")],
        ["현재 지식 상태", skills.get("knowledge_state")],
        ["다음 검토 트리거", skills.get("next_review_trigger")],
    ])

    doc.add_heading("7. 현재 상태", level=1)
    states = document.get("current_state", [])
    add_table(doc, ["항목", "확정", "구현·제작", "검증", "확인 필요", "보류", "책임 원본·증거"], normalize_table_rows(states, ["item", "confirmed", "implemented", "validated", "unresolved", "hold", "evidence"]) or [["[작성 필요]", "-", "-", "-", "-", "-", "-"]])

    doc.add_heading("8. 확정 결정", level=1)
    add_table(doc, ["ID", "결정", "이유", "영향 범위", "검증", "관련 스킬"], normalize_table_rows(document.get("decisions", []), ["id", "decision", "reason", "impact", "validation", "skills"]) or [["[작성 필요]", "-", "-", "-", "-", "-"]])

    doc.add_heading("9. 구현·제작 상태", level=1)
    add_table(doc, ["기능·자산", "실제 경로", "반영 범위", "남은 범위", "위험", "상태"], normalize_table_rows(document.get("implementation", []), ["item", "path", "scope", "remaining", "risk", "status"]) or [["[작성 필요]", "-", "-", "-", "-", "-"]])

    doc.add_heading("10. 검증 상태", level=1)
    add_table(doc, ["대상", "종류", "명령·방법", "결과", "증거", "미검증 이유"], normalize_table_rows(document.get("validation", []), ["target", "type", "method", "result", "evidence", "not_run_reason"]) or [["[작성 필요]", "-", "-", "-", "-", "-"]])

    for section_index, custom in enumerate(document.get("sections", []), start=11):
        doc.add_heading(f"{section_index}. {custom.get('title', '[상세 기획]')}", level=1)
        purpose = custom.get("purpose")
        if purpose:
            paragraph = doc.add_paragraph()
            style_run(paragraph.add_run(str(purpose)), 9)
        for block in custom.get("blocks", []):
            block_type = block.get("type")
            if block_type == "paragraph":
                paragraph = doc.add_paragraph()
                style_run(paragraph.add_run(_value(block.get("text"))), 9)
            elif block_type == "bullets":
                add_bullets(doc, block.get("items", []))
            elif block_type == "table":
                headers = [str(value) for value in block.get("headers", [])]
                rows = [[_value(value) for value in row] for row in block.get("rows", [])]
                if headers:
                    add_table(doc, headers, rows)

    doc.add_heading("승인 이미지·UI·다이어그램·실제 캡처", level=1)
    if approved_visuals:
        for index, (item, path) in enumerate(approved_visuals, start=1):
            add_image(doc, path, f"승인 자료 {index}. {item.get('title', item.get('asset_id', path.name))} - {item.get('caption', '')}", width=6.7)
            add_table(doc, ["Asset ID", "상태", "캐노니컬 경로", "채택 요소", "비채택 요소"], [[item.get("asset_id"), item.get("status"), item.get("path"), item.get("adopted_elements"), item.get("excluded_elements")]])
    else:
        doc.add_paragraph("[등록된 승인 이미지 없음]")

    doc.add_heading("위험·미확정·보류", level=1)
    add_table(doc, ["ID", "유형", "항목", "영향", "완화·재개 조건", "책임", "상태"], normalize_table_rows(document.get("risks", []), ["id", "type", "item", "impact", "mitigation_or_resume", "owner", "status"]) or [["[작성 필요]", "-", "-", "-", "-", "-", "-"]])

    doc.add_heading("다음 작업", level=1)
    add_table(doc, ["우선순위", "작업", "주 책임", "영향 분야", "선행 조건", "산출물", "완료 기준", "검증", "스킬"], normalize_table_rows(document.get("next_actions", []), ["priority", "task", "owner", "affected_disciplines", "prerequisites", "deliverables", "done", "validation", "skills"]) or [["1", "[작성 필요]", "-", "-", "-", "-", "-", "-", "-"]])

    doc.add_heading("Definition of Ready", level=1)
    add_bullets(doc, document.get("definition_of_ready", []))
    doc.add_heading("Definition of Done", level=1)
    add_bullets(doc, document.get("definition_of_done", []))

    doc.add_heading("부록·참조", level=1)
    add_table(doc, ["자료", "경로", "역할", "상태"], normalize_table_rows(document.get("appendices", []), ["title", "path", "role", "status"]) or [["[작성 필요]", "-", "-", "-"]])
    doc.add_heading("변경·학습 이력", level=1)
    add_table(doc, ["일자", "변경", "이유", "검증", "학습 상태"], normalize_table_rows(document.get("change_history", []), ["date", "change", "reason", "validation", "knowledge_state"]) or [["[작성 필요]", "-", "-", "-", "-"]])

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def convert_pdf(docx_path: Path, pdf_path: Path) -> None:
    executable = shutil.which("libreoffice") or shutil.which("soffice")
    if not executable:
        raise RuntimeError("LibreOffice is required to generate PDF outputs.")
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="design-doc-lo-") as temp:
        profile = Path(temp) / "profile"
        command = [
            executable,
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ]
        try:
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=120,
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError("LibreOffice PDF conversion timed out after 120 seconds.") from exc
        if result.returncode:
            raise RuntimeError(result.stdout + result.stderr)
    generated = pdf_path.parent / f"{docx_path.stem}.pdf"
    if generated != pdf_path:
        generated.replace(pdf_path)
    if not pdf_path.exists() or pdf_path.read_bytes()[:5] != b"%PDF-":
        raise RuntimeError(f"Invalid PDF output: {pdf_path}")


def render_pdf_for_review(pdf_path: Path, output_dir: Path) -> list[Path]:
    executable = shutil.which("pdftoppm")
    if not executable:
        raise RuntimeError("pdftoppm is required for automated render review.")
    output_dir.mkdir(parents=True, exist_ok=True)
    prefix = output_dir / "page"
    result = subprocess.run(
        [executable, "-png", "-r", "120", str(pdf_path), str(prefix)],
        capture_output=True,
        text=True,
    )
    if result.returncode:
        raise RuntimeError(result.stdout + result.stderr)
    pages = sorted(output_dir.glob("page-*.png"))
    if not pages:
        raise RuntimeError("PDF render review produced no pages.")
    for page in pages:
        with Image.open(page).convert("RGB") as image:
            if image.width < 500 or image.height < 500:
                raise RuntimeError(f"Rendered page is unexpectedly small: {page}")
            white = Image.new("RGB", image.size, "white")
            if ImageChops.difference(image, white).getbbox() is None:
                raise RuntimeError(f"Rendered page is blank: {page}")
    return pages


def approved_visuals(document: dict[str, Any], registry_dir: Path) -> list[tuple[dict[str, Any], Path]]:
    result: list[tuple[dict[str, Any], Path]] = []
    for item in document.get("approved_visuals", []):
        if not item.get("include_in_publication", True):
            continue
        path_value = str(item.get("path", "")).strip()
        if not path_value:
            continue
        path = resolve(registry_dir, path_value)
        if not path.is_file():
            raise FileNotFoundError(f"Approved visual missing: {path_value}")
        if path.suffix.lower() not in ALLOWED_VISUAL_SUFFIXES:
            raise ValueError(f"Unsupported approved visual format: {path_value}")
        result.append((item, path))
    return result


def build_one(entry: dict[str, Any], registry_path: Path, source_commit: str, human_review: str) -> dict[str, Any]:
    registry_dir = registry_path.parent
    source_path = resolve(registry_dir, str(entry["source_json"]))
    output_docx = resolve(registry_dir, str(entry["output_docx"]))
    output_pdf = resolve(registry_dir, str(entry["output_pdf"]))
    asset_dir = resolve(registry_dir, str(entry["asset_dir"]))
    manifest_path = resolve(registry_dir, str(entry["publication_manifest"]))
    generator_path = Path(__file__).resolve()
    diagram_module = generator_path.with_name("design_document_diagrams.py")

    document = load_json(source_path)
    if document.get("document_id") != entry.get("document_id"):
        raise ValueError(f"document_id mismatch: {source_path}")

    generated_dir = asset_dir / "generated"
    diagrams = {
        "workflow": generated_dir / "workflow.png",
        "status": generated_dir / "status-summary.png",
        "responsibility": generated_dir / "responsibility-map.png",
    }
    build_workflow_diagram(document, diagrams["workflow"])
    build_status_diagram(document, diagrams["status"])
    build_responsibility_diagram(document, diagrams["responsibility"])
    visuals = approved_visuals(document, registry_dir)
    source_hash = digest(source_path)
    build_docx(document, source_hash, source_commit, diagrams, visuals, output_docx)
    convert_pdf(output_docx, output_pdf)

    with tempfile.TemporaryDirectory(prefix="design-doc-render-") as temp:
        rendered_pages = render_pdf_for_review(output_pdf, Path(temp))
        page_count = len(rendered_pages)

    manifest = {
        "schema_version": 1,
        "publication_id": entry["document_id"],
        "role": "human-readable-derivative",
        "source_json": str(entry["source_json"]),
        "source_sha256": source_hash,
        "source_commit": source_commit,
        "generator": str(entry.get("generator", "tools/build_design_documents.py")),
        "generator_sha256": digest(generator_path),
        "diagram_generator_sha256": digest(diagram_module),
        "output_docx": str(entry["output_docx"]),
        "output_docx_sha256": digest(output_docx),
        "output_pdf": str(entry["output_pdf"]),
        "output_pdf_sha256": digest(output_pdf),
        "asset_dir": str(entry["asset_dir"]),
        "generated_diagrams": {os.path.relpath(path, registry_dir): digest(path) for path in diagrams.values()},
        "approved_visuals": {os.path.relpath(path, registry_dir): digest(path) for _, path in visuals},
        "generated_at": datetime.now(timezone.utc).replace(microsecond=0).isoformat(),
        "status": "CURRENT",
        "automated_render_review": "PASSED",
        "rendered_page_count": page_count,
        "human_visual_review": human_review,
        "editing_policy": "Edit the source JSON and regenerate. Do not maintain an active Markdown design bible.",
    }
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return manifest


def main() -> int:
    options = parse_args()
    registry_path = Path(options.registry).resolve()
    registry = load_json(registry_path)
    documents = registry.get("documents", [])
    selected = set(options.only)
    built = []
    for entry in documents:
        if entry.get("status") not in ACTIVE_STATUSES:
            continue
        if selected and entry.get("document_id") not in selected:
            continue
        built.append(build_one(entry, registry_path, options.source_commit, options.human_visual_review))
    if selected and len(built) != len(selected):
        built_ids = {item["publication_id"] for item in built}
        missing = sorted(selected - built_ids)
        raise SystemExit(f"Requested active document_id not found: {', '.join(missing)}")
    print(f"Generated {len(built)} design document publication(s).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
