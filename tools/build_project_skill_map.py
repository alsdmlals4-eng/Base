#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import json
import os
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from skill_map_diagrams import DISCIPLINES, LAYER_KO, STATUS_KO, discipline_diagram, flow_diagram, matrix_diagram

ACCENT, ACCENT_LIGHT, WHITE = "315EFB", "E8EEFF", "FFFFFF"


def args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate human-readable DOCX/PDF project skill maps from SKILL_REGISTRY.json")
    parser.add_argument("--registry", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--project-name", default="[프로젝트명]")
    parser.add_argument("--source-commit", default=os.environ.get("GITHUB_SHA", ""))
    parser.add_argument("--human-visual-review", default="NOT_RUN", choices=["NOT_RUN", "PASSED", "FAILED"])
    return parser.parse_args()


def digest(path: Path) -> str:
    hasher = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def load(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("Registry root must be an object")
    data.setdefault("routing_policy", {})
    data.setdefault("skills", [])
    data.setdefault("discipline_entrypoints", {})
    return data


def shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd")) or OxmlElement("w:shd")
    if node.getparent() is None:
        props.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell) -> None:
    props = cell._tc.get_or_add_tcPr()
    mar = props.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if mar.getparent() is None:
        props.append(mar)
    for name, value in (("top", 80), ("start", 100), ("bottom", 80), ("end", 100)):
        node = mar.find(qn(f"w:{name}")) or OxmlElement(f"w:{name}")
        if node.getparent() is None:
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def run_font(run, size=8, bold=False, color=None) -> None:
    run.font.name = "NanumGothic"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "NanumGothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    result = doc.add_table(rows=1, cols=len(headers))
    result.style = "Table Grid"
    result.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell = result.rows[0].cells[index]
        shade(cell, ACCENT)
        margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_font(cell.paragraphs[0].add_run(header), 8, True, (255, 255, 255))
    for row_index, values in enumerate(rows):
        cells = result.add_row().cells
        for index, value in enumerate(values):
            shade(cells[index], WHITE if row_index % 2 == 0 else "F4F6FA")
            margins(cells[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run_font(cells[index].paragraphs[0].add_run(value), 7.5)
    doc.add_paragraph()


def set_styles(doc: Document) -> None:
    for name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.name = "NanumGothic"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "NanumGothic")
    doc.styles["Normal"].font.size = Pt(9)
    doc.styles["Title"].font.size = Pt(30)
    doc.styles["Title"].font.bold = True
    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 1"].font.color.rgb = RGBColor(49, 94, 251)


def header_footer(doc: Document, project_name: str, registry_hash: str) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = f"{project_name} | PROJECT SKILL MAP"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run_font(run, 8, color=(102, 112, 133))
        footer = section.footer.paragraphs[0]
        footer.text = f"Registry SHA-256: {registry_hash[:16]}… | 자동 생성 파생본 - SKILL_REGISTRY.json이 책임 원본"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run_font(run, 7.5, color=(102, 112, 133))


def build_docx(registry: dict[str, Any], project_name: str, registry_hash: str, source_commit: str, assets: dict[str, Path], output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(1.7)
    set_styles(doc)
    header_footer(doc, project_name, registry_hash)

    title = doc.add_paragraph(style="Title")
    title.add_run(f"{project_name}\n프로젝트 스킬 지도")
    subtitle = doc.add_paragraph()
    run_font(subtitle.add_run("사람용 통합 열람본 · 이미지 포함 · SKILL_REGISTRY.json에서 자동 생성"), 11, color=(102, 112, 133))

    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Table Grid"
    pairs = [("책임 원본", "SKILL_REGISTRY.json"), ("생성 역할", "읽기 전용 파생본"), ("Registry 해시", registry_hash[:20] + "…"), ("기준 커밋", source_commit or "[작성 필요]")]
    for index, (label, value) in enumerate(pairs):
        row, pair = divmod(index, 2)
        label_cell, value_cell = meta.rows[row].cells[pair * 2], meta.rows[row].cells[pair * 2 + 1]
        shade(label_cell, ACCENT_LIGHT)
        margins(label_cell); margins(value_cell)
        run_font(label_cell.paragraphs[0].add_run(label), 9, True)
        run_font(value_cell.paragraphs[0].add_run(value), 9)
    doc.add_paragraph()

    policy = registry["routing_policy"]
    active = [item for item in registry["skills"] if item.get("status") in ("ACTIVE", "SUPPORT")]
    coverage = sum(1 for discipline in DISCIPLINES if registry["discipline_entrypoints"].get(discipline))
    doc.add_heading("1. 한눈에 보기", level=1)
    table(doc, ["항목", "현재 기준"], [
        ["전체 스킬 자동 로드", "금지" if not policy.get("load_all_skills", False) else "허용"],
        ["기본 선택", str(policy.get("default_selection", "none"))], ["현행·보조 스킬", str(len(active))],
        ["분야 진입 스킬 등록", f"{coverage}/{len(DISCIPLINES)}"], ["주 책임 분야 스킬 최대", str(policy.get("max_primary_discipline_skills", 1))],
        ["Foundation 스킬 최대", str(policy.get("max_foundation_skills", 3))],
    ])
    doc.add_picture(str(assets["flow"]), width=Inches(7.1))
    caption = doc.add_paragraph("그림 1. 요청부터 Learning Log까지의 선택적 호출 흐름"); caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("2. 분야별 진입 구조", level=1)
    doc.add_picture(str(assets["discipline"]), width=Inches(7.1))
    caption = doc.add_paragraph("그림 2. 11개 책임 분야의 진입 스킬 등록 상태"); caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table(doc, ["분야", "진입 스킬", "상태"], [[discipline, ", ".join(registry["discipline_entrypoints"].get(discipline, [])) or "[설치 필요]", "등록" if registry["discipline_entrypoints"].get(discipline) else "미등록"] for discipline in DISCIPLINES])

    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = Cm(1.4)
    section.left_margin = section.right_margin = Cm(1.2)
    doc.add_heading("3. 현행 스킬 매트릭스", level=1)
    doc.add_picture(str(assets["matrix"]), width=Inches(9.7))
    caption = doc.add_paragraph("그림 3. Registry 기반 스킬 상태와 호출 조건 요약"); caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    skills = registry["skills"] or [registry.get("skill_contract_example", {})]
    table(doc, ["스킬", "계층", "분야", "상태", "Trigger", "사용 조건", "비호출 조건", "지식 상태"], [[
        str(item.get("skill_id", "[스킬 등록 필요]")), LAYER_KO.get(str(item.get("layer", "")), str(item.get("layer", ""))), str(item.get("discipline", "-")),
        STATUS_KO.get(str(item.get("status", "NOT_INSTALLED")), str(item.get("status", ""))), " / ".join(item.get("trigger_tags", [])) or "-",
        "\n".join(item.get("use_when", [])[:2]) or "-", "\n".join(item.get("do_not_use_when", [])[:2]) or "-", str(item.get("knowledge_state", "OBSERVATION")),
    ] for item in skills])

    doc.add_heading("4. 선택적 호출 규칙", level=1)
    for text in ["전체 skills 폴더를 기본 컨텍스트로 읽지 않는다.", "trigger가 일치하지 않는 스킬은 호출하지 않는다.", "주 책임 분야 스킬은 동시에 최대 1개만 선택한다.", "Foundation 스킬은 현재 절차에 필요한 최소 개수만 선택한다.", "PDF 발행·Handoff·Health Review는 해당 단계에서만 호출한다.", "보류·백업·제거 후보 스킬은 재개 승인 전 호출하지 않는다."]:
        doc.add_paragraph(text, style="List Bullet")
    doc.add_heading("5. 학습·갱신 계약", level=1)
    table(doc, ["대상", "계약"], [
        ["모든 의미 있는 호출", "결과·실패·예외·사용자 피드백·과다 호출·누락 검증을 Learning Log에 기록"],
        ["스킬 본문 갱신", "반복 실패, 새 예외, 책임·경로·검증 변경처럼 근거가 있을 때만 수행"],
        ["지식 상태", "관찰 → 가설 → 패턴 → 검증 → 승격 후보"], ["공용화 경계", "프로젝트 고유 수치·세계관·승인 자산은 프로젝트에 유지"],
    ])
    doc.add_heading("6. 운영·검수 체크", level=1)
    for text in ["Registry의 현행 스킬 경로가 실제 존재한다.", "각 활성 분야의 진입 스킬 또는 통합 책임이 등록돼 있다.", "DOCX/PDF와 다이어그램은 Registry 해시와 일치한다.", "사람은 이 문서/PDF를 보고, AI는 Registry만 읽는다.", "스킬 변경 시 Registry·Learning Log·DOCX/PDF를 같은 작업에서 갱신한다.", "문서가 생성됐다는 사실만으로 실제 스킬 실행·검증 완료로 표시하지 않는다."]:
        doc.add_paragraph(text, style="List Bullet")
    doc.save(output)


def convert_pdf(docx: Path, output_dir: Path) -> Path:
    executable = shutil.which("libreoffice") or shutil.which("soffice")
    if not executable:
        raise RuntimeError("LibreOffice is required to generate PROJECT_SKILL_MAP.pdf")
    with tempfile.TemporaryDirectory(prefix="skill-map-lo-") as temp:
        profile = Path(temp) / "profile"
        profile.mkdir()
        command = [
            executable,
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx),
        ]
        try:
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=120,
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError("LibreOffice skill-map conversion timed out after 120 seconds.") from exc
        if result.returncode:
            raise RuntimeError(result.stdout + result.stderr)
    pdf = output_dir / f"{docx.stem}.pdf"
    if not pdf.exists() or pdf.read_bytes()[:5] != b"%PDF-":
        raise RuntimeError("LibreOffice did not create a valid PDF")
    return pdf


def main() -> int:
    options = args()
    registry_path, output_dir = Path(options.registry).resolve(), Path(options.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    assets_dir = output_dir / "PROJECT_SKILL_MAP.assets"; assets_dir.mkdir(exist_ok=True)
    registry, registry_hash = load(registry_path), digest(registry_path)
    assets = {"flow": assets_dir / "skill-flow.png", "discipline": assets_dir / "discipline-routing.png", "matrix": assets_dir / "skill-matrix.png"}
    flow_diagram(assets["flow"]); discipline_diagram(registry, assets["discipline"]); matrix_diagram(registry, assets["matrix"])
    docx = output_dir / "PROJECT_SKILL_MAP.docx"
    build_docx(registry, options.project_name, registry_hash, options.source_commit, assets, docx)
    pdf = convert_pdf(docx, output_dir)
    manifest = {
        "schema_version": 1, "publication_id": "project-skill-map", "role": "human-readable-derivative",
        "source_registry": registry_path.name if registry_path.parent == output_dir else str(registry_path), "source_sha256": registry_hash,
        "source_commit": options.source_commit, "output_docx": docx.name, "output_docx_sha256": digest(docx),
        "output_pdf": pdf.name, "output_pdf_sha256": digest(pdf), "diagram_paths": [str(path.relative_to(output_dir)) for path in assets.values()],
        "diagram_sha256": {str(path.relative_to(output_dir)): digest(path) for path in assets.values()}, "generator": "tools/build_project_skill_map.py",
        "generated_at": datetime.now(timezone.utc).replace(microsecond=0).isoformat(), "status": "CURRENT",
        "automated_render_review": "PASSED", "human_visual_review": options.human_visual_review,
        "editing_policy": "Edit SKILL_REGISTRY.json and regenerate. Do not maintain PROJECT_SKILL_MAP.md.",
    }
    (output_dir / "SKILL_MAP_PUBLICATION_MANIFEST.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"Generated {docx} and {pdf}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
